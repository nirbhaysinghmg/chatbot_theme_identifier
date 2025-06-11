import os
os.environ["PROTOCOL_BUFFERS_PYTHON_IMPLEMENTATION"] = "python"

import uuid
from fastapi import FastAPI, HTTPException, WebSocket, WebSocketDisconnect, Body, UploadFile, File, BackgroundTasks, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_chroma import Chroma
from langchain.text_splitter import CharacterTextSplitter
from langchain.chains import ConversationalRetrievalChain
from langchain_core.documents import Document
import pytesseract
from PIL import Image
import io
import tempfile
import shutil
from typing import List, Optional, Dict, Any
from langchain.prompts import PromptTemplate
import json
import time
from datetime import datetime
import PyPDF2
import docx
import asyncio
from concurrent.futures import ThreadPoolExecutor
import logging
from queue import Queue
import threading
import signal
import sys

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment
from dotenv import load_dotenv
load_dotenv()

# Initialize FastAPI with WebSocket support
app = FastAPI(title="Document Analysis Chatbot")

# Add CORS middleware with specific origins
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Add startup and shutdown events
import httpx

@app.on_event("startup")
async def startup_event():
    print("Application startup...")
    # Start keep-alive background task
    async def keep_alive():
        while True:
            try:
                async with httpx.AsyncClient() as client:
                    await client.get("http://localhost:8008/health")
            except Exception as e:
                print(f"Keep-alive ping failed: {e}")
            await asyncio.sleep(600)  # 10 minutes
    asyncio.create_task(keep_alive())

@app.on_event("shutdown")
async def shutdown_event():
    print("Application shutdown...")
    # Clean up any resources here
    for session_id in list(document_collections.keys()):
        try:
            session_dir = os.path.join(PERSIST_DIRECTORY, session_id)
            if os.path.exists(session_dir):
                shutil.rmtree(session_dir)
        except Exception as e:
            print(f"Error cleaning up session {session_id}: {str(e)}")

# Prepare embeddings
EMBED_MODEL = os.getenv("EMBED_MODEL", "models/embedding-001")
embeddings = GoogleGenerativeAIEmbeddings(model=EMBED_MODEL, google_api_key=os.getenv("GEMINI_API_KEY"))

# Initialize ChromaDB vector store
PERSIST_DIRECTORY = os.getenv("PERSIST_DIRECTORY", "chroma_db")

TESSERACT_PATH = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

def get_vector_store(session_id=None):
    # Create a unique directory for each session
    if session_id:
        persist_dir = os.path.join(PERSIST_DIRECTORY, session_id)
        # Create directory if it doesn't exist
        os.makedirs(persist_dir, exist_ok=True)
    else:
        persist_dir = PERSIST_DIRECTORY
    
    return Chroma(
        persist_directory=persist_dir,
        embedding_function=embeddings
    )

# Initialize LLM
llm = ChatGoogleGenerativeAI(
    model=os.getenv("LLM_MODEL", "gemini-2.0-flash"),
    google_api_key=os.getenv("GEMINI_API_KEY"),
    disable_streaming=True
)

# Define system prompt for theme analysis
THEME_ANALYSIS_PROMPT = PromptTemplate(
    input_variables=["context", "question", "chat_history"],
    template="""You are a document analysis assistant. Your role is to:
    Talk to the user in a friendly and engaging manner, and also provide a direct answer to the user's question based on the document excerpts, If the users greet you, greet them back in a friendly manner and ask them what they want to know about the document excerpts
    If the user ask you something that is not in the document they have uploaded, tell them that you are not sure about the answer and ask them if they want to know something else
1. Analyze the provided document excerpts
2. Identify main themes and patterns
3. Provide a synthesized summary with citations

Format your response exactly like this:
First line of the response should be the answer to the user's question based on the document excerpts, and then the rest of the response should be the themes and the supporting document citations

Themes:
[For each identified theme, provide:
- Theme title
- Summary of the theme
- Supporting document citations (DocID, Page, Paragraph)]

Context: {context}
Chat History: {chat_history}
Question: {question}

Remember to:
1. Keep theme summaries clear and concise
2. Include specific citations for each theme
3. Group related information under appropriate themes
4. Format the response in Markdown
"""
)

# Store chat histories and document collections for different sessions
chat_histories = {}
document_collections = {}

# Add a new dictionary to track original files
original_files = {}

# Document processing queue
document_queue = Queue()
processing_status = {}

# Thread pool for document processing
thread_pool = ThreadPoolExecutor(max_workers=4)

# Pydantic models
class QueryRequest(BaseModel):
    question: str
    session_id: Optional[str] = None

class DocumentMetadata(BaseModel):
    title: str
    description: Optional[str] = None
    tags: Optional[List[str]] = None
    category: Optional[str] = None

class BulkUploadResponse(BaseModel):
    success: List[str]
    failed: List[dict]

class BulkUploadRequest(BaseModel):
    files: List[UploadFile]
    metadata: Optional[Dict[str, Any]] = None
    session_id: Optional[str] = None

class ProcessingStatus(BaseModel):
    status: str
    progress: float
    total_files: int
    processed_files: int
    failed_files: List[str]
    completed_files: List[str]

# Document processing functions
async def process_pdf(file: UploadFile) -> list:
    contents = await file.read()
    pdf_file = io.BytesIO(contents)
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    page_texts = []
    for i, page in enumerate(pdf_reader.pages):
        page_text = page.extract_text()
        if page_text:
            page_texts.append({"text": page_text, "page": i + 1})
    return page_texts

async def process_docx(file: UploadFile) -> list:
    contents = await file.read()
    docx_file = io.BytesIO(contents)
    doc = docx.Document(docx_file)
    para_texts = []
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            para_texts.append({"text": paragraph.text, "paragraph": i + 1})
    return para_texts

async def process_txt(file: UploadFile) -> str:
    contents = await file.read()
    return contents.decode('utf-8')

async def process_image_with_ocr(file: UploadFile) -> str:
    try:
        contents = await file.read()
        image = Image.open(io.BytesIO(contents))
        text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing image: {str(e)}")

@app.post("/query")
async def query_qa(req: QueryRequest):
    try:
        # Generate a session ID if not provided
        session_id = req.session_id or str(uuid.uuid4())
        
        # Check if session has documents
        if session_id not in document_collections or not document_collections[session_id]:
            raise HTTPException(
                status_code=400,
                detail="No documents available for analysis. Please upload documents first."
            )
        
        # Initialize chat history if needed
        if session_id not in chat_histories:
            chat_histories[session_id] = []
        
        # Get vector store for this session
        vector_store = document_collections[session_id]
        
        # Build retriever with proper search parameters
        retriever = vector_store.as_retriever(
            search_type="similarity",
            search_kwargs={"k": 5}  # Number of documents to retrieve
        )
        
        # Use ConversationalRetrievalChain with source documents
        qa = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=retriever,
            return_source_documents=True,
            combine_docs_chain_kwargs={"prompt": THEME_ANALYSIS_PROMPT},
            verbose=True
        )
        
        # Get answer using chat history
        result = qa({
            "question": req.question,
            "chat_history": chat_histories[session_id]
        })
        
        # Process source documents
        source_docs = result.get("source_documents", [])
        individual_answers = []
        
        for doc in source_docs:
            page_num = doc.metadata.get("page", "Unknown")
            para_num = doc.metadata.get("paragraph", "Unknown")
            chunk_index = doc.metadata.get("chunk_index", "Unknown")
            
            doc_answer = {
                "doc_id": doc.metadata.get("source", "Unknown"),
                "answer": doc.page_content,
                "citation": {
                    "page": page_num,
                    "paragraph": para_num,
                    "chunk": chunk_index
                }
            }
            individual_answers.append(doc_answer)
        
        # Update chat history
        chat_histories[session_id].append((req.question, result["answer"]))
        
        # Limit chat history length
        if len(chat_histories[session_id]) > 10:
            chat_histories[session_id] = chat_histories[session_id][-10:]
        
        return {
            "answer": result["answer"],
            "individual_answers": individual_answers,
            "themes": extract_themes(result["answer"]),
            "session_id": session_id  # Return the session ID for client reference
        }
    except HTTPException as he:
        raise he
    except Exception as e:
        print(f"Error processing query: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error processing query: {str(e)}"
        )

def extract_themes(answer: str) -> List[dict]:
    """Extract themes from the LLM response."""
    themes = []
    current_theme = None
    
    for line in answer.split('\n'):
        if line.startswith('Theme'):
            if current_theme:
                themes.append(current_theme)
            current_theme = {
                'title': line.split(':', 1)[1].strip(),
                'summary': '',
                'citations': []
            }
        elif current_theme and line.strip():
            if '(' in line and ')' in line:
                # Extract citation
                citation_text = line[line.find('(')+1:line.find(')')]
                if 'Page' in citation_text and 'Para' in citation_text:
                    page = citation_text.split('Page')[1].split(',')[0].strip()
                    para = citation_text.split('Para')[1].strip()
                    current_theme['citations'].append({
                        'page': page,
                        'paragraph': para
                    })
            else:
                current_theme['summary'] += line.strip() + ' '
    
    if current_theme:
        themes.append(current_theme)
    
    return themes

@app.post("/upload/document")
async def upload_document(
    file: UploadFile = File(...),
    metadata: Optional[dict] = None,
    session_id: str = None
):
    print(f"Uploading document for session: {session_id}")
    session_id = session_id or str(uuid.uuid4())
    print(f"Using session ID: {session_id}")
    
    try:
        # Validate file type
        file_extension = file.filename.split('.')[-1].lower()
        allowed_extensions = ['pdf', 'docx', 'doc', 'txt', 'jpg', 'jpeg', 'png']
        
        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type. Allowed types: {', '.join(allowed_extensions)}"
            )

        # Process document based on file type
        if file_extension == 'pdf':
            page_texts = await process_pdf(file)
            # Split each page into paragraphs and keep track of (page, paragraph, text)
            pdf_paragraphs = []
            for p in page_texts:
                # Split on double newlines or single newlines as fallback
                paras = [para.strip() for para in p["text"].split('\n\n') if para.strip()]
                if len(paras) == 1:
                    # If only one paragraph, try splitting on single newlines
                    paras = [para.strip() for para in p["text"].split('\n') if para.strip()]
                for idx, para in enumerate(paras):
                    pdf_paragraphs.append({"text": para, "page": p["page"], "paragraph": idx + 1})
            text = "\n".join([para["text"] for para in pdf_paragraphs])
        elif file_extension in ['docx', 'doc']:
            para_texts = await process_docx(file)
            text = "\n".join([p["text"] for p in para_texts])
        elif file_extension == 'txt':
            text = await process_txt(file)
        elif file_extension in ['jpg', 'jpeg', 'png']:
            text = await process_image_with_ocr(file)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")

        if not text.strip():
            raise HTTPException(
                status_code=400,
                detail="No text content could be extracted from the document"
            )

        # Create document with metadata
        doc = Document(
            page_content=text,
            metadata={
                "source": file.filename,
                "type": file_extension,
                "upload_time": datetime.now().isoformat(),
                "title": file.filename.rsplit('.', 1)[0],  # Remove extension for title
                **(metadata or {})
            }
        )
        
        # Split text into chunks
        splitter = CharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separator="\n"
        )
        chunks = splitter.split_text(text)
        
        if not chunks:
            raise HTTPException(
                status_code=400,
                detail="Document could not be split into meaningful chunks"
            )
        
        # Create documents for each chunk
        documents = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = doc.metadata.copy()
            chunk_metadata.update({
                "chunk_index": i,
                "total_chunks": len(chunks),
                "is_chunk": True,
                "original_file": file.filename
            })
            # Assign page/paragraph if possible
            if file_extension == 'pdf':
                # Find which paragraph (and page) this chunk belongs to
                para_num = None
                page_num = None
                char_count = 0
                for p in pdf_paragraphs:
                    if char_count + len(p["text"]) >= sum(len(c) for c in chunks[:i+1]):
                        page_num = p["page"]
                        para_num = p["paragraph"]
                        break
                    char_count += len(p["text"])
                if page_num:
                    chunk_metadata["page"] = page_num
                if para_num:
                    chunk_metadata["paragraph"] = para_num
            elif file_extension in ['docx', 'doc']:
                para_num = None
                char_count = 0
                for p in para_texts:
                    if char_count + len(p["text"]) >= sum(len(c) for c in chunks[:i+1]):
                        para_num = p["paragraph"]
                        break
                    char_count += len(p["text"])
                if para_num:
                    chunk_metadata["paragraph"] = para_num
            documents.append(Document(page_content=chunk, metadata=chunk_metadata))
        
        # Generate unique IDs for documents
        ids = [str(uuid.uuid4()) for _ in documents]
        
        # Initialize or get vector store for this session
        if session_id not in document_collections:
            print(f"Creating new document collection for session {session_id}")
            document_collections[session_id] = get_vector_store(session_id)
            original_files[session_id] = []  # Initialize original files list for this session
        
        # Add to vector store
        print(f"Adding {len(documents)} chunks to vector store for session {session_id}")
        document_collections[session_id].add_documents(documents=documents, ids=ids)
        
        # Add to original files list
        original_files[session_id].append({
            "title": file.filename,
            "type": file_extension,
            "upload_time": datetime.now().isoformat(),
            "description": metadata.get('description') if metadata else None,
            "chunks": len(chunks)
        })
        
        return {
            "status": "success",
            "message": f"Successfully uploaded {file.filename}",
            "chunks": len(chunks),
            "document_id": ids[0],  # Return the first document ID as reference
            "session_id": session_id  # Return the session ID
        }
    except HTTPException as he:
        raise he
    except Exception as e:
        print(f"Error processing document: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error processing document: {str(e)}"
        )

@app.get("/documents")
async def list_documents(session_id: str = None):
    print(f"Listing documents for session: {session_id}")
    
    # If no session ID provided, try to get from query params
    if not session_id:
        session_id = "default"
    
    if session_id not in document_collections:
        print(f"No document collection found for session {session_id}")
        return {"documents": []}
    
    try:
        # Return the original files list instead of chunks
        documents = original_files.get(session_id, [])
        print(f"Found {len(documents)} original files for session {session_id}")
        return {"documents": documents}
    except Exception as e:
        print(f"Error listing documents: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/documents/{document_id}")
async def delete_document(document_id: str, session_id: str = None):
    session_id = session_id or "default"
    
    if session_id not in document_collections:
        raise HTTPException(status_code=404, detail="No documents found for this session")
    
    try:
        document_collections[session_id].delete(ids=[document_id])
        return {"status": "success", "message": f"Document {document_id} deleted successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/session/{session_id}")
async def clear_session(session_id: str):
    if session_id in chat_histories:
        del chat_histories[session_id]
    if session_id in document_collections:
        del document_collections[session_id]
    if session_id in original_files:
        del original_files[session_id]
    # Clean up the session's ChromaDB directory
    session_dir = os.path.join(PERSIST_DIRECTORY, session_id)
    if os.path.exists(session_dir):
        shutil.rmtree(session_dir)
    return {"status": "success", "message": f"Session {session_id} cleared successfully"}

@app.delete("/sessions/clear")
async def clear_all_sessions():
    chat_histories.clear()
    document_collections.clear()
    original_files.clear()
    # Clean up all session directories
    if os.path.exists(PERSIST_DIRECTORY):
        for item in os.listdir(PERSIST_DIRECTORY):
            item_path = os.path.join(PERSIST_DIRECTORY, item)
            if os.path.isdir(item_path):
                shutil.rmtree(item_path)
    return {"status": "success", "message": "All sessions cleared successfully"}

# Add WebSocket connection manager
class ConnectionManager:
    def __init__(self):
        self.active_connections: Dict[str, WebSocket] = {}

    async def connect(self, websocket: WebSocket, session_id: str):
        await websocket.accept()
        self.active_connections[session_id] = websocket

    def disconnect(self, session_id: str):
        if session_id in self.active_connections:
            del self.active_connections[session_id]

    async def send_message(self, message: dict, session_id: str):
        if session_id in self.active_connections:
            await self.active_connections[session_id].send_json(message)

manager = ConnectionManager()

@app.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    session_id = None
    try:
        # Create a unique session ID for this WebSocket connection
        session_id = str(uuid.uuid4())
        print(f"New WebSocket connection attempt for session {session_id}")
        
        await manager.connect(websocket, session_id)
        print(f"WebSocket connection accepted for session {session_id}")
        
        # Initialize session data
        if session_id not in chat_histories:
            chat_histories[session_id] = []
        if session_id not in document_collections:
            document_collections[session_id] = get_vector_store(session_id)
        
        while True:
            try:
                # Receive message from client
                data = await websocket.receive_text()
                print(f"Received message from client session {session_id}: {data[:100]}...")
                message = json.loads(data)
                
                if "user_input" in message:
                    # Check if session has documents
                    if not document_collections[session_id]:
                        await manager.send_message({
                            "error": "No documents available for analysis. Please upload documents first.",
                            "done": True
                        }, session_id)
                        continue
                    
                    # Get chat history from message
                    chat_history = message.get("chat_history", [])
                    if chat_history:
                        formatted_history = [(msg["content"], "") for msg in chat_history if msg["role"] == "user"]
                        chat_histories[session_id] = formatted_history
                    
                    # Build retriever
                    retriever = document_collections[session_id].as_retriever(
                        search_type="similarity",
                        search_kwargs={"k": 5}
                    )
                    
                    # Use ConversationalRetrievalChain with source documents
                    qa = ConversationalRetrievalChain.from_llm(
                        llm=llm,
                        retriever=retriever,
                        return_source_documents=True,
                        combine_docs_chain_kwargs={"prompt": THEME_ANALYSIS_PROMPT},
                        verbose=True
                    )
                    
                    try:
                        # Get answer using chat history
                        result = qa({
                            "question": message["user_input"],
                            "chat_history": chat_histories[session_id]
                        })
                        
                        # Process source documents
                        source_docs = result.get("source_documents", [])
                        individual_answers = []
                        
                        for doc in source_docs:
                            page_num = doc.metadata.get("page", "Unknown")
                            para_num = doc.metadata.get("paragraph", "Unknown")
                            chunk_index = doc.metadata.get("chunk_index", "Unknown")
                            
                            doc_answer = {
                                "doc_id": doc.metadata.get("source", "Unknown"),
                                "answer": doc.page_content,
                                "citation": {
                                    "page": page_num,
                                    "paragraph": para_num,
                                    "chunk": chunk_index
                                }
                            }
                            individual_answers.append(doc_answer)
                        
                        # Update chat history
                        chat_histories[session_id].append((message["user_input"], result["answer"]))
                        
                        # Send response back to client
                        response = {
                            "text": result["answer"],
                            "individual_answers": individual_answers,
                            "themes": extract_themes(result["answer"]),
                            "done": True
                        }
                        await manager.send_message(response, session_id)
                        print(f"Response sent successfully for session {session_id}")
                        
                    except Exception as e:
                        error_msg = f"Error processing request: {str(e)}"
                        print(error_msg)
                        await manager.send_message({
                            "error": error_msg,
                            "done": True
                        }, session_id)
                        
            except WebSocketDisconnect:
                print(f"WebSocket disconnected for session {session_id}")
                manager.disconnect(session_id)
                break
                
    except Exception as e:
        print(f"Fatal WebSocket error: {str(e)}")
    finally:
        if session_id:
            print(f"Cleaning up session {session_id}")
            # Clean up session data
            if session_id in chat_histories:
                del chat_histories[session_id]
            if session_id in document_collections:
                del document_collections[session_id]
            if session_id in original_files:
                del original_files[session_id]
            # Clean up the session's ChromaDB directory
            session_dir = os.path.join(PERSIST_DIRECTORY, session_id)
            if os.path.exists(session_dir):
                shutil.rmtree(session_dir)
            manager.disconnect(session_id)
        try:
            await websocket.close()
        except:
            pass

# Keep the original endpoint for backward compatibility
@app.websocket("/ws/chat")
async def websocket_endpoint_chat(websocket: WebSocket):
    await websocket_endpoint(websocket)

# Add health check endpoints
@app.get("/")
async def root():
    return {"status": "healthy", "message": "API is running"}

@app.head("/")
async def head_root():
    return JSONResponse(content=None, status_code=200)

@app.get("/health")
async def health_check():
    return {"status": "healthy"}

async def process_document_in_background(file: UploadFile, session_id: str, metadata: Optional[Dict] = None):
    try:
        # Process document based on file type
        file_extension = file.filename.split('.')[-1].lower()
        
        if file_extension == 'pdf':
            page_texts = await process_pdf(file)
            # Split each page into paragraphs and keep track of (page, paragraph, text)
            pdf_paragraphs = []
            for p in page_texts:
                paras = [para.strip() for para in p["text"].split('\n\n') if para.strip()]
                if len(paras) == 1:
                    paras = [para.strip() for para in p["text"].split('\n') if para.strip()]
                for idx, para in enumerate(paras):
                    pdf_paragraphs.append({"text": para, "page": p["page"], "paragraph": idx + 1})
            text = "\n".join([para["text"] for para in pdf_paragraphs])
        elif file_extension in ['docx', 'doc']:
            para_texts = await process_docx(file)
            text = "\n".join([p["text"] for p in para_texts])
        elif file_extension == 'txt':
            text = await process_txt(file)
        elif file_extension in ['jpg', 'jpeg', 'png']:
            text = await process_image_with_ocr(file)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

        if not text.strip():
            raise ValueError("No text content could be extracted from the document")

        # Create document with metadata
        doc = Document(
            page_content=text,
            metadata={
                "source": file.filename,
                "type": file_extension,
                "upload_time": datetime.now().isoformat(),
                "title": file.filename.rsplit('.', 1)[0],
                **(metadata or {})
            }
        )
        
        # Split text into chunks
        splitter = CharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separator="\n"
        )
        chunks = splitter.split_text(text)
        
        if not chunks:
            raise ValueError("Document could not be split into meaningful chunks")
        
        # Create documents for each chunk
        documents = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = doc.metadata.copy()
            chunk_metadata.update({
                "chunk_index": i,
                "total_chunks": len(chunks),
                "is_chunk": True,
                "original_file": file.filename
            })
            # Assign page/paragraph if possible
            if file_extension == 'pdf':
                para_num = None
                page_num = None
                char_count = 0
                for p in pdf_paragraphs:
                    if char_count + len(p["text"]) >= sum(len(c) for c in chunks[:i+1]):
                        page_num = p["page"]
                        para_num = p["paragraph"]
                        break
                    char_count += len(p["text"])
                if page_num:
                    chunk_metadata["page"] = page_num
                if para_num:
                    chunk_metadata["paragraph"] = para_num
            elif file_extension in ['docx', 'doc']:
                para_num = None
                char_count = 0
                for p in para_texts:
                    if char_count + len(p["text"]) >= sum(len(c) for c in chunks[:i+1]):
                        para_num = p["paragraph"]
                        break
                    char_count += len(p["text"])
                if para_num:
                    chunk_metadata["paragraph"] = para_num
            documents.append(Document(page_content=chunk, metadata=chunk_metadata))
        
        # Generate unique IDs for documents
        ids = [str(uuid.uuid4()) for _ in documents]
        
        # Add to vector store
        document_collections[session_id].add_documents(documents=documents, ids=ids)
        
        # Update processing status
        if session_id in processing_status:
            processing_status[session_id]["processed_files"] += 1
            processing_status[session_id]["completed_files"].append(file.filename)
            processing_status[session_id]["progress"] = (
                processing_status[session_id]["processed_files"] / 
                processing_status[session_id]["total_files"] * 100
            )
        
        return True, file.filename
        
    except Exception as e:
        logger.error(f"Error processing document {file.filename}: {str(e)}")
        if session_id in processing_status:
            processing_status[session_id]["failed_files"].append(file.filename)
        return False, file.filename

@app.post("/upload/bulk")
async def bulk_upload_documents(
    background_tasks: BackgroundTasks,
    request: BulkUploadRequest
):
    session_id = request.session_id or str(uuid.uuid4())
    
    # Initialize processing status
    processing_status[session_id] = {
        "status": "processing",
        "progress": 0,
        "total_files": len(request.files),
        "processed_files": 0,
        "failed_files": [],
        "completed_files": []
    }
    
    # Initialize vector store for session if needed
    if session_id not in document_collections:
        document_collections[session_id] = get_vector_store(session_id)
        original_files[session_id] = []
    
    # Process documents in background
    background_tasks.add_task(
        process_bulk_documents,
        request.files,
        session_id,
        request.metadata
    )
    
    return {
        "status": "processing",
        "session_id": session_id,
        "message": f"Processing {len(request.files)} documents"
    }

async def process_bulk_documents(
    files: List[UploadFile],
    session_id: str,
    metadata: Optional[Dict] = None
):
    try:
        # Process documents concurrently
        tasks = []
        for file in files:
            task = asyncio.create_task(
                process_document_in_background(file, session_id, metadata)
            )
            tasks.append(task)
        
        # Wait for all tasks to complete
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Update final status
        processing_status[session_id]["status"] = "completed"
        
    except Exception as e:
        logger.error(f"Error in bulk processing: {str(e)}")
        processing_status[session_id]["status"] = "failed"
        processing_status[session_id]["error"] = str(e)

@app.get("/upload/status/{session_id}")
async def get_upload_status(session_id: str):
    if session_id not in processing_status:
        raise HTTPException(status_code=404, detail="Session not found")
    
    return processing_status[session_id]

def handle_sigterm(signum, frame):
    print("Received SIGTERM. Starting graceful shutdown...")
    sys.exit(0)

if __name__ == "__main__":
    import uvicorn
    
    # Register signal handlers
    signal.signal(signal.SIGTERM, handle_sigterm)
    
    # Get port from environment variable or use default
    port = int(os.getenv("PORT", 8008))
    
    # Configure uvicorn
    config = uvicorn.Config(
        app=app,
        host="0.0.0.0",
        port=port,
        workers=1,  # Use single worker for stability
        timeout_keep_alive=30,  # Reduce keep-alive timeout
        log_level="info"
    )
    
    print(f"\nStarting server on http://0.0.0.0:{port}")
    server = uvicorn.Server(config)
    server.run()














